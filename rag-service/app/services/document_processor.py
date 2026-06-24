import io
import boto3
from botocore.client import Config
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import structlog
from app.config import get_settings

logger = structlog.get_logger(__name__)
settings = get_settings()


def _get_s3_client():
    return boto3.client(
        "s3",
        endpoint_url=settings.do_spaces_endpoint,
        aws_access_key_id=settings.do_spaces_key,
        aws_secret_access_key=settings.do_spaces_secret,
        config=Config(signature_version="s3v4"),
        region_name="us-east-1",
    )


def download_document(storage_key: str) -> bytes:
    """Download a document from DigitalOcean Spaces."""
    s3 = _get_s3_client()
    logger.info("Downloading document", storage_key=storage_key)

    response = s3.get_object(
        Bucket=settings.do_spaces_bucket,
        Key=storage_key,
    )
    content = response["Body"].read()
    logger.info("Document downloaded", size_bytes=len(content))
    return content


def extract_text_from_pdf(content: bytes) -> str:
    """Extract clean text from PDF using PyMuPDF."""
    doc = fitz.open(stream=content, filetype="pdf")
    pages: list[str] = []

    for page_num, page in enumerate(doc):
        text = page.get_text("text")  # type: ignore[call-arg]
        if text.strip():
            # Clean up common PDF artifacts
            cleaned = (
                text
                .replace("\x00", "")          # null bytes
                .replace("\ufeff", "")         # BOM
                .strip()
            )
            pages.append(f"[Page {page_num + 1}]\n{cleaned}")

    doc.close()
    full_text = "\n\n".join(pages)
    logger.info("PDF text extracted", page_count=len(pages), char_count=len(full_text))
    return full_text


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX including tables."""
    doc = DocxDocument(io.BytesIO(content))
    sections: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            sections.append(para.text.strip())

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                sections.append(" | ".join(row_texts))

    full_text = "\n\n".join(sections)
    logger.info("DOCX text extracted", section_count=len(sections), char_count=len(full_text))
    return full_text


def extract_text_from_txt(content: bytes) -> str:
    """Decode plain text with encoding fallback."""
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return content.decode("utf-8", errors="replace")


def extract_text(storage_key: str, mime_type: str) -> str:
    """Download and extract text from a document."""
    content = download_document(storage_key)

    if mime_type == "application/pdf":
        return extract_text_from_pdf(content)
    elif mime_type in (
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ):
        return extract_text_from_docx(content)
    elif mime_type == "text/plain":
        return extract_text_from_txt(content)
    else:
        raise ValueError(f"Unsupported MIME type for text extraction: {mime_type}")
